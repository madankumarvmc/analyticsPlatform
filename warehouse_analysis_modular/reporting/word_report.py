#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Microsoft Word Report Generator

Generates professional MS Word documents for warehouse analysis reports
with AI-powered insights and embedded charts.
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import pandas as pd

# Import from parent directories
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from config import REPORT_DIR, CHARTS_DIR
from warehouse_analysis_modular.reporting.llm_integration import LLMIntegration
from warehouse_analysis_modular.reporting.prompts_config import get_prompt_by_type, get_all_prompts
from warehouse_analysis_modular.utils.helpers import setup_logging

logger = setup_logging()


class WordReportGenerator:
    """
    Generates professional Microsoft Word reports with AI insights.
    """
    
    def __init__(self, 
                 output_dir: Path = REPORT_DIR,
                 charts_dir: Path = CHARTS_DIR,
                 llm_integration: Optional[LLMIntegration] = None):
        """
        Initialize the Word report generator.
        
        Args:
            output_dir: Directory for output files
            charts_dir: Directory containing chart images
            llm_integration: LLM integration instance for AI insights
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required for Word report generation. Install with: pip install python-docx")
        
        self.output_dir = Path(output_dir)
        self.charts_dir = Path(charts_dir)
        self.llm_integration = llm_integration or LLMIntegration()
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # Ensure output directory exists and test write permissions
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            # Test write permissions by creating a temporary file
            test_file = self.output_dir / "test_write_permissions.tmp"
            test_file.write_text("test")
            test_file.unlink()  # Remove test file
            self.logger.info(f"Output directory confirmed writable: {self.output_dir}")
        except Exception as e:
            self.logger.error(f"Cannot write to output directory {self.output_dir}: {e}")
            # Fallback to current working directory
            import tempfile
            self.output_dir = Path(tempfile.gettempdir()) / "warehouse_reports"
            self.output_dir.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Using fallback directory: {self.output_dir}")
    
    def _add_heading_with_style(self, doc: Document, text: str, level: int = 1):
        """Add a styled heading to the document."""
        heading = doc.add_heading(text, level=level)
        if level == 1:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True
        elif level == 2:
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True
        return heading
    
    def _add_paragraph_with_style(self, doc: Document, text: str, style: str = None):
        """Add a styled paragraph to the document."""
        paragraph = doc.add_paragraph(text)
        if style:
            paragraph.style = style
        return paragraph
    
    def _add_ai_insight_section(self, doc: Document, title: str, insight_text: str):
        """Add an AI insight section with special formatting and bullet points."""
        # Add insight heading
        heading = doc.add_paragraph()
        run = heading.add_run(f"💡 AI Insights: {title}")
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Parse and format bullet points with bold headers
        insight_lines = insight_text.split('\n')
        for line in insight_lines:
            line = line.strip()
            if not line:
                continue
                
            insight_para = doc.add_paragraph()
            
            # Check if line starts with bullet point
            if line.startswith('•') or line.startswith('-'):
                # Remove bullet marker and split on bold markers
                clean_line = line.lstrip('•-').strip()
                
                # Look for bold headers (text between ** markers)
                if '**' in clean_line:
                    parts = clean_line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Odd indices are bold text
                            run = insight_para.add_run(part)
                            run.font.bold = True
                        else:  # Even indices are regular text
                            insight_para.add_run(part)
                else:
                    # No bold formatting found
                    insight_para.add_run(clean_line)
                    
                # Apply bullet style
                insight_para.style = 'List Bullet'
            else:
                # Regular paragraph
                insight_para.add_run(line)
                insight_para.style = 'Quote'
        
        return heading
    
    def _add_chart_with_insights(self, doc: Document, chart_name: str, chart_path: Path, analysis_results: Dict):
        """Add a chart with AI-generated insights."""
        # Add chart title regardless of whether image is available
        chart_title = chart_name.replace('_', ' ').title()
        self._add_heading_with_style(doc, chart_title, level=3)
        
        # Try to add chart image
        chart_added = False
        if chart_path.exists():
            try:
                self.logger.info(f"Adding chart image: {chart_path}")
                doc.add_picture(str(chart_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_added = True
                self.logger.info(f"Chart {chart_name} added successfully")
            except Exception as e:
                self.logger.error(f"Failed to add chart image {chart_name}: {e}")
        else:
            self.logger.warning(f"Chart image not found: {chart_path}")
        
        # Add placeholder if chart couldn't be added
        if not chart_added:
            placeholder = doc.add_paragraph(f"📊 Chart: {chart_title}")
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.style = 'Intense Quote'
            note = doc.add_paragraph(f"Note: Chart image ({chart_name}) not available in this report version.")
            note.style = 'Caption'
        
        # Generate AI insights for this chart (even without the image)
        try:
            chart_insights = self._generate_chart_insights(chart_name, analysis_results)
            if chart_insights and not chart_insights.startswith("("):
                self._add_ai_insight_section(doc, f"{chart_title} Analysis", chart_insights)
        except Exception as e:
            self.logger.error(f"Failed to generate insights for chart {chart_name}: {e}")
        
        # Add spacing
        doc.add_paragraph()
    
    def _generate_chart_insights(self, chart_name: str, analysis_results: Dict) -> str:
        """Generate AI insights for a specific chart."""
        # Get prompt for this chart
        prompt_config = get_prompt_by_type('chart', chart_name)
        
        if not prompt_config.get('instruction'):
            return ""
        
        # Build context facts based on chart type
        facts = self._build_chart_context_facts(chart_name, analysis_results)
        
        # Generate prompt
        prompt = self.llm_integration.build_prompt(
            prompt_config['context'], 
            facts, 
            prompt_config['instruction']
        )
        
        # Get AI response
        return self.llm_integration.call_gemini(prompt)
    
    def _build_chart_context_facts(self, chart_name: str, analysis_results: Dict) -> Dict[str, Any]:
        """Build context facts for chart analysis."""
        facts = {"Chart type": chart_name}
        
        # Add relevant data based on chart type
        if 'date' in chart_name:
            date_summary = analysis_results.get('date_order_summary')
            if date_summary is not None and not date_summary.empty:
                facts.update({
                    "Date range": f"{date_summary['Date'].min().date()} to {date_summary['Date'].max().date()}",
                    "Peak volume": f"{date_summary['Total_Case_Equiv'].max():.0f}",
                    "Average volume": f"{date_summary['Total_Case_Equiv'].mean():.0f}",
                    "Volume variation": f"{date_summary['Total_Case_Equiv'].std():.0f}"
                })
        
        elif 'sku' in chart_name:
            sku_summary = analysis_results.get('sku_order_summary')
            if sku_summary is not None and not sku_summary.empty:
                facts.update({
                    "Total SKUs": len(sku_summary),
                    "Top SKU volume": f"{sku_summary.iloc[0].get('Order_Volume_CE', 0):.0f}" if len(sku_summary) > 0 else "N/A"
                })
        
        elif 'abc' in chart_name:
            abc_summary = analysis_results.get('abc_fms_summary')
            if abc_summary is not None and not abc_summary.empty:
                facts.update({
                    "ABC classes": len(abc_summary['ABC'].unique()) if 'ABC' in abc_summary.columns else "N/A"
                })
        
        return facts
    
    def _add_data_table(self, doc: Document, title: str, data: pd.DataFrame, max_rows: int = 3):
        """Add a formatted data table to the document."""
        if data is None or data.empty:
            return
        
        # Add table title
        self._add_heading_with_style(doc, title, level=3)
        
        # Limit rows
        display_data = data.head(max_rows)
        
        # Create table
        table = doc.add_table(rows=1, cols=len(display_data.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, column in enumerate(display_data.columns):
            header_cells[i].text = str(column)
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for _, row in display_data.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if pd.isna(value):
                    row_cells[i].text = "N/A"
                elif isinstance(value, (int, float)):
                    if isinstance(value, float) and value != int(value):
                        row_cells[i].text = f"{value:.2f}"
                    else:
                        row_cells[i].text = f"{int(value):,}"
                else:
                    row_cells[i].text = str(value)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_executive_summary(self, doc: Document, analysis_results: Dict):
        """Add executive summary with enhanced data extraction (exactly 2 paragraphs)."""
        self._add_heading_with_style(doc, "Executive Summary", level=1)
        
        # Extract specific numbers for enhanced executive summary
        stats = analysis_results.get('order_statistics', {})
        date_summary = analysis_results.get('date_order_summary')
        abc_summary = analysis_results.get('abc_fms_summary')
        
        # Calculate enhanced metrics
        enhanced_facts = {
            "Total dates analyzed": stats.get('unique_dates', 0),
            "Total SKUs": stats.get('unique_skus', 0),
            "Total order lines": stats.get('total_order_lines', 0),
            "Total case equivalents": stats.get('total_case_equivalent', 0),
            "Date range": "N/A",
            "Peak vs average ratio": "N/A",
            "ABC classification breakdown": "N/A"
        }
        
        # Calculate date range
        if date_summary is not None and not date_summary.empty and 'Date' in date_summary.columns:
            date_range_start = date_summary['Date'].min().strftime('%Y-%m-%d')
            date_range_end = date_summary['Date'].max().strftime('%Y-%m-%d')
            enhanced_facts["Date range"] = f"{date_range_start} to {date_range_end}"
            
            # Calculate peak vs average ratio
            if 'Total_Case_Equiv' in date_summary.columns:
                peak_volume = date_summary['Total_Case_Equiv'].max()
                avg_volume = date_summary['Total_Case_Equiv'].mean()
                if avg_volume > 0:
                    enhanced_facts["Peak vs average ratio"] = f"{peak_volume/avg_volume:.1f}x"
        
        # Calculate ABC breakdown percentages
        if abc_summary is not None and not abc_summary.empty:
            if 'ABC' in abc_summary.columns:
                abc_counts = abc_summary['ABC'].value_counts()
                total_skus = len(abc_summary)
                if total_skus > 0:
                    a_pct = (abc_counts.get('A', 0) / total_skus) * 100
                    b_pct = (abc_counts.get('B', 0) / total_skus) * 100
                    c_pct = (abc_counts.get('C', 0) / total_skus) * 100
                    enhanced_facts["ABC classification breakdown"] = f"A:{a_pct:.0f}%, B:{b_pct:.0f}%, C:{c_pct:.0f}%"
        
        # Generate AI-powered executive summary with enhanced data
        try:
            executive_summary = self.llm_integration.generate_cover_summary(analysis_results, enhanced_facts)
            if executive_summary and not executive_summary.startswith("("):
                # Split into exactly 2 paragraphs and format
                paragraphs = [p.strip() for p in executive_summary.split('\n\n') if p.strip()]
                for i, para in enumerate(paragraphs[:2]):  # Limit to exactly 2 paragraphs
                    self._add_paragraph_with_style(doc, para)
                    if i == 0:  # Add spacing between paragraphs
                        doc.add_paragraph()
            else:
                # Fallback summary with specific numbers
                # Fallback executive summary as bullet points
                fallback_summary = f"""• **Data Scope**: {enhanced_facts['Total dates analyzed']} days, {enhanced_facts['Total SKUs']:,} SKUs, {enhanced_facts['Total order lines']:,} order lines analyzed
• **Volume Scale**: {enhanced_facts['Total case equivalents']:,.0f} case equivalents processed from {enhanced_facts['Date range']}
• **Demand Pattern**: Peak-to-average ratio of {enhanced_facts['Peak vs average ratio']} indicates operational variability
• **Classification**: ABC distribution shows {enhanced_facts['ABC classification breakdown']} requiring strategic focus"""
                self._add_ai_insight_section(doc, "Summary", fallback_summary)
        except Exception as e:
            self.logger.error(f"Failed to generate executive summary: {e}")
        
        doc.add_page_break()
    
    def _add_key_findings(self, doc: Document, analysis_results: Dict):
        """Add key findings section (limited to exactly 1 page)."""
        self._add_heading_with_style(doc, "Key Findings", level=1)
        
        # Generate AI-powered key findings with enhanced data
        try:
            prompt_config = get_prompt_by_type('word', 'key_findings_summary')
            stats = analysis_results.get('order_statistics', {})
            date_summary = analysis_results.get('date_order_summary')
            sku_summary = analysis_results.get('sku_order_summary')
            abc_summary = analysis_results.get('abc_fms_summary')
            
            # Calculate specific metrics for key findings
            enhanced_facts = {
                "Analysis period": f"{stats.get('unique_dates', 0)} days",
                "SKU diversity": f"{stats.get('unique_skus', 0):,} unique SKUs",
                "Order complexity": f"{stats.get('total_order_lines', 0):,} order lines",
                "Volume processed": f"{stats.get('total_case_equivalent', 0):,.0f} case equivalents"
            }
            
            # Add calculated ratios and percentages
            if date_summary is not None and not date_summary.empty and 'Total_Case_Equiv' in date_summary.columns:
                peak_volume = date_summary['Total_Case_Equiv'].max()
                avg_volume = date_summary['Total_Case_Equiv'].mean()
                if avg_volume > 0:
                    enhanced_facts["Peak to average ratio"] = f"{peak_volume/avg_volume:.1f}x"
                    enhanced_facts["Volume variability"] = f"{(date_summary['Total_Case_Equiv'].std()/avg_volume)*100:.0f}%"
            
            # Add Pareto analysis if available
            if sku_summary is not None and not sku_summary.empty:
                total_volume = sku_summary['Order_Volume_CE'].sum()
                cumulative_pct = (sku_summary['Order_Volume_CE'].cumsum() / total_volume) * 100
                sku_count_80pct = len(cumulative_pct[cumulative_pct <= 80])
                enhanced_facts["Pareto ratio"] = f"{(sku_count_80pct/len(sku_summary))*100:.0f}% of SKUs drive 80% of volume"
            
            prompt = self.llm_integration.build_prompt(
                prompt_config['context'], 
                enhanced_facts, 
                prompt_config['instruction']
            )
            
            key_findings = self.llm_integration.call_gemini(prompt)
            if key_findings and not key_findings.startswith("("):
                # Format as bullet points with bold headers
                self._add_ai_insight_section(doc, "Summary", key_findings)
            else:
                # Fallback findings with ultra-concise format
                fallback_findings = f"""• **Volume Concentration**: {enhanced_facts.get('Pareto ratio', 'N/A')}
• **Demand Variability**: Peak demand is **{enhanced_facts.get('Peak to average ratio', 'N/A')}** higher than average
• **Operational Scale**: **{enhanced_facts['Volume processed']}** across **{enhanced_facts['SKU diversity']}**"""
                self._add_ai_insight_section(doc, "Summary", fallback_findings)
        except Exception as e:
            self.logger.error(f"Failed to generate key findings: {e}")
        
        doc.add_page_break()
    
    def generate_word_report(self, analysis_results: Dict, filename: str = None) -> Path:
        """
        Generate a comprehensive Word report.
        
        Args:
            analysis_results: Dictionary containing all analysis results
            filename: Output filename (auto-generated if None)
            
        Returns:
            Path to the generated Word document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required")
        
        self.logger.info("Starting Word report generation")
        
        # Detect environment (helpful for debugging deployment issues)
        import os
        env_info = {
            "platform": os.name,
            "cwd": os.getcwd(),
            "streamlit_detected": "streamlit" in os.environ.get("PATH", "").lower() or "STREAMLIT_SERVER_PORT" in os.environ,
            "python_path": sys.executable if hasattr(sys, 'executable') else "Unknown"
        }
        self.logger.info(f"Environment info: {env_info}")
        
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Warehouse_Analysis_Report_{timestamp}.docx"
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Warehouse Analysis Report"
        doc.core_properties.subject = "Operational Analysis and Recommendations"
        doc.core_properties.author = "Warehouse Analysis Tool"
        doc.core_properties.created = datetime.now()
        
        # Add title page
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("WAREHOUSE ANALYSIS REPORT")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Operational Insights and Strategic Recommendations")
        subtitle_run.font.size = Pt(14)
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        
        doc.add_page_break()
        
        # Add executive summary
        self._add_executive_summary(doc, analysis_results)
        
        # Add key findings
        self._add_key_findings(doc, analysis_results)
        
        # Add merged daily operations analysis (combines volume and customer patterns)
        self._add_heading_with_style(doc, "Daily Operations Analysis", level=1)
        
        # Generate merged date profile summary (combines volume + customer insights)
        try:
            merged_insights = self.llm_integration.generate_date_profile_merged_summary(analysis_results)
            if merged_insights and not merged_insights.startswith("("):
                self._add_ai_insight_section(doc, "Daily Demand Patterns (Volume & Customer)", merged_insights)
        except Exception as e:
            self.logger.error(f"Failed to generate merged date insights: {e}")
        
        # Add date summary table (limited to top 5 days to save space)
        date_summary = analysis_results.get('date_order_summary')
        if date_summary is not None and not date_summary.empty:
            # Format date column for display
            display_date_summary = date_summary.copy()
            if 'Date' in display_date_summary.columns:
                display_date_summary['Date'] = display_date_summary['Date'].dt.strftime('%Y-%m-%d')
            self._add_data_table(doc, "Daily Operations Summary (Top 3 Peak Days)", 
                               display_date_summary.nlargest(3, 'Total_Case_Equiv'), max_rows=3)
        
        # Add enhanced multi-line order trend chart (replaces separate volume and customer charts)
        enhanced_trend_chart_path = self.charts_dir / 'enhanced_order_trend_profile.png'
        
        # Enhanced debugging for chart paths
        self.logger.info(f"Looking for enhanced trend chart at: {enhanced_trend_chart_path}")
        self.logger.info(f"Charts directory contents: {list(self.charts_dir.glob('*.png')) if self.charts_dir.exists() else 'Directory does not exist'}")
        
        if enhanced_trend_chart_path.exists():
            file_size = enhanced_trend_chart_path.stat().st_size
            self.logger.info(f"✅ Found enhanced order trend chart: {enhanced_trend_chart_path} (size: {file_size} bytes)")
            self._add_chart_with_insights(doc, 'enhanced_order_trend_profile', enhanced_trend_chart_path, analysis_results)
        else:
            self.logger.warning(f"❌ Enhanced order trend chart not found: {enhanced_trend_chart_path}")
            self.logger.info("Falling back to original separate volume and customer charts")
            
            # Fallback to original charts if enhanced chart is not available
            date_chart_path = self.charts_dir / 'date_total_case_equiv.png'
            customer_chart_path = self.charts_dir / 'date_distinct_customers.png'
            
            # Add volume chart with compact insights
            if date_chart_path.exists():
                self._add_chart_with_insights(doc, 'date_total_case_equiv', date_chart_path, analysis_results)
            
            # Add customer chart with compact insights  
            if customer_chart_path.exists():
                self._add_chart_with_insights(doc, 'date_distinct_customers', customer_chart_path, analysis_results)
        
        doc.add_page_break()
        
        # Add capacity planning section
        self._add_heading_with_style(doc, "Capacity Planning Analysis", level=1)
        
        # Percentile insights
        try:
            percentile_insights = self.llm_integration.generate_percentile_summary(analysis_results)
            if percentile_insights and not percentile_insights.startswith("("):
                self._add_ai_insight_section(doc, "Capacity Requirements", percentile_insights)
        except Exception as e:
            self.logger.error(f"Failed to generate percentile insights: {e}")
        
        # Add percentile table (reduced size)
        percentile_data = analysis_results.get('percentile_profile')
        if percentile_data is not None and not percentile_data.empty:
            # Show only key percentiles
            key_percentiles = percentile_data[percentile_data['Percentile'].isin(['50%ile', '75%ile', '95%ile', 'Max'])]
            self._add_data_table(doc, "Key Volume Percentiles", key_percentiles, max_rows=4)
        
        # Add percentile chart
        percentile_chart_path = self.charts_dir / 'percentile_total_case_equiv.png'
        if percentile_chart_path.exists():
            self._add_chart_with_insights(doc, 'percentile_total_case_equiv', percentile_chart_path, analysis_results)
        
        doc.add_page_break()
        
        # Add SKU analysis section (limited to 2-3 pages max)
        self._add_heading_with_style(doc, "SKU Distribution Analysis", level=1)
        
        # SKU insights (concise format)
        try:
            sku_insights = self.llm_integration.generate_sku_profile_summary(analysis_results)
            if sku_insights and not sku_insights.startswith("("):
                self._add_ai_insight_section(doc, "SKU Performance Patterns", sku_insights)
        except Exception as e:
            self.logger.error(f"Failed to generate SKU insights: {e}")
        
        # Add SKU summary table (reduced to top 3 to save space)
        sku_summary = analysis_results.get('sku_order_summary')
        if sku_summary is not None and not sku_summary.empty:
            self._add_data_table(doc, "Top SKUs by Volume", sku_summary.head(3), max_rows=3)
        
        # Add SKU Pareto chart with concise insights
        sku_pareto_path = self.charts_dir / 'sku_pareto.png'
        if sku_pareto_path.exists():
            self._add_chart_with_insights(doc, 'sku_pareto', sku_pareto_path, analysis_results)
        
        doc.add_page_break()
        
        # Add ABC-FMS classification section
        self._add_heading_with_style(doc, "ABC-FMS Classification Analysis", level=1)
        
        # ABC-FMS insights
        try:
            abc_fms_insights = self.llm_integration.generate_abc_fms_summary(analysis_results)
            if abc_fms_insights and not abc_fms_insights.startswith("("):
                self._add_ai_insight_section(doc, "Classification Strategy", abc_fms_insights)
        except Exception as e:
            self.logger.error(f"Failed to generate ABC-FMS insights: {e}")
        
        # Add ABC-FMS summary table (reduced size)
        abc_fms_summary = analysis_results.get('abc_fms_summary')
        if abc_fms_summary is not None and not abc_fms_summary.empty:
            # Show only key combinations
            key_combinations = abc_fms_summary.head(3)
            self._add_data_table(doc, "Key ABC-FMS Classifications", key_combinations, max_rows=3)
        
        # Add ABC volume chart
        abc_volume_path = self.charts_dir / 'abc_volume_stacked.png'
        if abc_volume_path.exists():
            self._add_chart_with_insights(doc, 'abc_volume_stacked', abc_volume_path, analysis_results)
        
        # Add ABC-FMS heatmap
        abc_heatmap_path = self.charts_dir / 'abc_fms_heatmap.png'
        if abc_heatmap_path.exists():
            self._add_chart_with_insights(doc, 'abc_fms_heatmap', abc_heatmap_path, analysis_results)
        
        # Add enhanced SKU Profile 2D Classification chart
        sku_2d_chart_path = self.charts_dir / 'sku_profile_2d_classification.png'
        
        # Enhanced debugging for 2D chart paths
        self.logger.info(f"Looking for SKU 2D classification chart at: {sku_2d_chart_path}")
        
        if sku_2d_chart_path.exists():
            file_size = sku_2d_chart_path.stat().st_size
            self.logger.info(f"✅ Found SKU 2D classification chart: {sku_2d_chart_path} (size: {file_size} bytes)")
            self._add_chart_with_insights(doc, 'sku_profile_2d_classification', sku_2d_chart_path, analysis_results)
        else:
            self.logger.warning(f"❌ SKU 2D classification chart not found: {sku_2d_chart_path}")
            # Add placeholder text in Word report to indicate missing chart
            placeholder = doc.add_paragraph("📊 Enhanced SKU Profile 2D Classification Chart")
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.style = 'Intense Quote'
            note = doc.add_paragraph("Note: Enhanced 2D classification chart is being generated and will be available in future reports.")
            note.style = 'Caption'
        
        doc.add_page_break()
        
        # Add advanced analysis sections if available
        self._add_advanced_analysis_sections(doc, analysis_results)
        
        # Add recommendations section
        self._add_heading_with_style(doc, "Strategic Recommendations", level=1)
        
        try:
            prompt_config = get_prompt_by_type('word', 'recommendations_summary')
            facts = {
                "Analysis scope": f"{analysis_results.get('order_statistics', {}).get('unique_dates', 'N/A')} days analyzed",
                "Operational complexity": f"{analysis_results.get('order_statistics', {}).get('unique_skus', 'N/A')} SKUs managed",
                "Volume scale": f"{analysis_results.get('order_statistics', {}).get('total_case_equivalent', 0):.0f} case equivalents"
            }
            
            prompt = self.llm_integration.build_prompt(
                prompt_config['context'], 
                facts, 
                prompt_config['instruction']
            )
            
            recommendations = self.llm_integration.call_gemini(prompt)
            if recommendations and not recommendations.startswith("("):
                self._add_paragraph_with_style(doc, recommendations)
            else:
                # Fallback recommendations as bullets
                fallback_recommendations = """• **Immediate**: Focus on top **20%** of SKUs driving **80%** of volume for optimization
• **Medium-term**: Implement **6-month** capacity planning based on **95th percentile** requirements  
• **Strategic**: Develop ABC-based slotting strategy for **long-term** operational efficiency"""
                self._add_ai_insight_section(doc, "Action Plan", fallback_recommendations)
        except Exception as e:
            self.logger.error(f"Failed to generate recommendations: {e}")
        
        # Add advanced analysis sections if available
        self._add_advanced_analysis_sections(doc, analysis_results)
        
        # Save document with enhanced error handling
        output_path = self.output_dir / filename
        
        try:
            self.logger.info(f"Saving Word document to: {output_path}")
            doc.save(str(output_path))
            
            # Verify file was created successfully
            if not output_path.exists():
                raise FileNotFoundError(f"Document was not created at expected path: {output_path}")
            
            file_size = output_path.stat().st_size
            if file_size == 0:
                raise ValueError("Generated document is empty (0 bytes)")
            
            self.logger.info(f"Word report generated successfully: {output_path} ({file_size:,} bytes)")
            return output_path
            
        except PermissionError as e:
            self.logger.error(f"Permission denied writing to {output_path}: {e}")
            # Try alternative location
            import tempfile
            alt_path = Path(tempfile.gettempdir()) / filename
            self.logger.info(f"Attempting to save to alternative location: {alt_path}")
            doc.save(str(alt_path))
            self.logger.info(f"Word report saved to alternative location: {alt_path}")
            return alt_path
            
        except Exception as e:
            self.logger.error(f"Failed to save Word document: {e}")
            raise
    
    def _add_advanced_analysis_sections(self, doc: Document, analysis_results: Dict):
        """Add advanced analysis sections if available."""
        
        # Multi-Metric Correlation Analysis
        if 'advanced_order_analysis' in analysis_results:
            advanced_analysis = analysis_results['advanced_order_analysis']
            
            if 'correlation_analysis' in advanced_analysis:
                self._add_heading_with_style(doc, "Multi-Metric Correlation Analysis", level=1)
                
                try:
                    correlation_insights = self.llm_integration.generate_correlation_summary(
                        advanced_analysis['correlation_analysis']
                    )
                    if correlation_insights and not correlation_insights.startswith("("):
                        self._add_ai_insight_section(doc, "Operational Metrics Correlation", correlation_insights)
                except Exception as e:
                    self.logger.error(f"Failed to generate correlation insights: {e}")
                
                # Add correlation matrix table if available
                if 'correlation_matrix' in advanced_analysis['correlation_analysis']:
                    corr_matrix = advanced_analysis['correlation_analysis']['correlation_matrix']
                    if not corr_matrix.empty:
                        self._add_data_table(doc, "Correlation Matrix", corr_matrix, max_rows=5)
                
                doc.add_page_break()
        
        # Picking Methodology Analysis
        if 'picking_analysis' in analysis_results:
            self._add_heading_with_style(doc, "Case vs Piece Picking Analysis", level=1)
            
            try:
                picking_insights = self.llm_integration.generate_picking_summary(
                    analysis_results['picking_analysis']
                )
                if picking_insights and not picking_insights.startswith("("):
                    self._add_ai_insight_section(doc, "Picking Methodology Insights", picking_insights)
            except Exception as e:
                self.logger.error(f"Failed to generate picking insights: {e}")
            
            # Add picking summary table
            picking_analysis = analysis_results['picking_analysis']
            if 'overall_picking_patterns' in picking_analysis:
                picking_summary = picking_analysis['overall_picking_patterns'].get('picking_summary')
                if picking_summary is not None and not picking_summary.empty:
                    self._add_data_table(doc, "Picking Method Distribution", picking_summary, max_rows=5)
            
            doc.add_page_break()
        
        # Enhanced ABC-FMS 2D Classification
        if 'enhanced_abc_fms_analysis' in analysis_results:
            self._add_heading_with_style(doc, "2D Classification Matrix Analysis", level=1)
            
            try:
                enhanced_abc_insights = self.llm_integration.generate_enhanced_abc_fms_summary(
                    analysis_results['enhanced_abc_fms_analysis']
                )
                if enhanced_abc_insights and not enhanced_abc_insights.startswith("("):
                    self._add_ai_insight_section(doc, "2D Matrix Classification", enhanced_abc_insights)
            except Exception as e:
                self.logger.error(f"Failed to generate enhanced ABC-FMS insights: {e}")
            
            # Add classification matrix summary
            enhanced_analysis = analysis_results['enhanced_abc_fms_analysis']
            if 'classification_matrix_2d' in enhanced_analysis:
                matrices = enhanced_analysis['classification_matrix_2d']
                if 'combined_matrix' in matrices:
                    combined_matrix = matrices['combined_matrix']
                    if not combined_matrix.empty:
                        self._add_data_table(doc, "2D Classification Summary", combined_matrix, max_rows=8)
            
            doc.add_page_break()
        
        # Advanced Capacity Planning
        if 'advanced_order_analysis' in analysis_results:
            advanced_analysis = analysis_results['advanced_order_analysis']
            
            if 'enhanced_percentile_analysis' in advanced_analysis and 'peak_ratios' in advanced_analysis:
                self._add_heading_with_style(doc, "Advanced Capacity Planning", level=1)
                
                try:
                    capacity_insights = self.llm_integration.generate_advanced_capacity_summary(
                        advanced_analysis['enhanced_percentile_analysis'],
                        advanced_analysis['peak_ratios']
                    )
                    if capacity_insights and not capacity_insights.startswith("("):
                        self._add_ai_insight_section(doc, "Capacity Planning Strategy", capacity_insights)
                except Exception as e:
                    self.logger.error(f"Failed to generate advanced capacity insights: {e}")
                
                # Add capacity recommendations table
                if 'capacity_recommendations' in advanced_analysis:
                    capacity_recs = advanced_analysis['capacity_recommendations']
                    if 'volume_capacity' in capacity_recs:
                        vol_capacity = capacity_recs['volume_capacity']
                        
                        # Create capacity planning summary table
                        capacity_data = []
                        capacity_data.append(['Design Capacity Base', f"{vol_capacity.get('design_capacity_base', 0):,.0f}"])
                        capacity_data.append(['Recommended Buffer', f"{vol_capacity.get('recommended_buffer_percentage', 0):.0f}%"])
                        capacity_data.append(['Total Design Capacity', f"{vol_capacity.get('total_design_capacity', 0):,.0f}"])
                        capacity_data.append(['Avg Utilization', f"{vol_capacity.get('utilization_at_avg', 0):.0f}%"])
                        
                        capacity_df = pd.DataFrame(capacity_data, columns=['Metric', 'Value'])
                        self._add_data_table(doc, "Capacity Planning Recommendations", capacity_df, max_rows=4)
                
                doc.add_page_break()
        
        # Operational Complexity Assessment
        if 'advanced_order_analysis' in analysis_results:
            advanced_analysis = analysis_results['advanced_order_analysis']
            
            if 'operational_complexity' in advanced_analysis:
                self._add_heading_with_style(doc, "Operational Complexity Assessment", level=1)
                
                try:
                    complexity_insights = self.llm_integration.generate_complexity_summary(
                        advanced_analysis['operational_complexity']
                    )
                    if complexity_insights and not complexity_insights.startswith("("):
                        self._add_ai_insight_section(doc, "Complexity Analysis", complexity_insights)
                except Exception as e:
                    self.logger.error(f"Failed to generate complexity insights: {e}")
                
                # Add complexity factors table
                complexity_data = advanced_analysis['operational_complexity']
                if 'complexity_factors' in complexity_data:
                    factors = complexity_data['complexity_factors']
                    
                    complexity_table_data = []
                    complexity_table_data.append(['Overall Score', f"{complexity_data.get('overall_complexity_score', 0):.1f}"])
                    complexity_table_data.append(['Complexity Level', complexity_data.get('complexity_level', 'Unknown')])
                    complexity_table_data.append(['Volume Variability', f"{factors.get('volume_variability', 0):.1f}"])
                    complexity_table_data.append(['Multi-truck Complexity', f"{factors.get('multi_truck_complexity', 0):.1f}%"])
                    
                    complexity_df = pd.DataFrame(complexity_table_data, columns=['Factor', 'Value'])
                    self._add_data_table(doc, "Complexity Factors", complexity_df, max_rows=4)
                
                doc.add_page_break()
        
        # Category Performance Distribution Analysis
        if 'category_performance_analysis' in analysis_results:
            self._add_heading_with_style(doc, "Category Performance Distribution Analysis", level=1)
            
            category_analysis = analysis_results['category_performance_analysis']
            
            try:
                # Generate AI insights for category performance
                category_insights = self.llm_integration.generate_category_performance_summary(
                    category_analysis
                )
                if category_insights and not category_insights.startswith("("):
                    self._add_ai_insight_section(doc, "Strategic Slotting Insights", category_insights)
            except Exception as e:
                self.logger.error(f"Failed to generate category performance insights: {e}")
            
            # Add performance summary table
            if 'performance_summary' in category_analysis:
                performance_summary = category_analysis['performance_summary']
                if not performance_summary.empty:
                    # Show top 10 categories by priority score
                    top_categories = performance_summary.head(10)
                    self._add_data_table(doc, "Category Performance Summary", top_categories, max_rows=10)
            
            # Add slotting recommendations
            if 'slotting_insights' in category_analysis:
                insights = category_analysis['slotting_insights']
                if 'slotting_recommendations' in insights:
                    recommendations = insights['slotting_recommendations']
                    
                    # Create recommendations summary
                    rec_summary_data = []
                    
                    # High priority recommendations
                    dock_proximity = recommendations.get('dock_proximity', [])
                    if dock_proximity:
                        rec_summary_data.append([
                            'Dock Proximity', 
                            ', '.join(dock_proximity[:3]), 
                            'High volume & velocity'
                        ])
                    
                    # Medium priority recommendations  
                    structured_bins = recommendations.get('structured_bins', [])
                    if structured_bins:
                        rec_summary_data.append([
                            'Structured Bins',
                            ', '.join(structured_bins[:5]),
                            'Moderate activity level'
                        ])
                    
                    # Standard storage
                    standard_storage = recommendations.get('standard_storage', [])
                    if standard_storage and len(standard_storage) <= 5:
                        rec_summary_data.append([
                            'Standard Storage',
                            ', '.join(standard_storage),
                            'Low activity level'
                        ])
                    elif len(standard_storage) > 5:
                        rec_summary_data.append([
                            'Standard Storage',
                            f'{len(standard_storage)} categories',
                            'Low activity level'
                        ])
                    
                    if rec_summary_data:
                        rec_df = pd.DataFrame(rec_summary_data, columns=['Strategy', 'Categories', 'Rationale'])
                        self._add_data_table(doc, "Slotting Strategy Recommendations", rec_df, max_rows=3)
            
            # Add the three distribution tables
            
            # SKU Distribution Table
            if 'sku_distribution' in category_analysis:
                sku_dist = category_analysis['sku_distribution']
                if not sku_dist.empty:
                    # Show top categories only to keep table manageable
                    sku_dist_top = sku_dist.head(8)  # Show top 7 categories + Grand Total
                    self._add_data_table(doc, "SKU % Distribution by Category", sku_dist_top, max_rows=8)
            
            # Cases Distribution Table
            if 'cases_distribution' in category_analysis:
                cases_dist = category_analysis['cases_distribution']
                if not cases_dist.empty:
                    # Show top categories only
                    cases_dist_top = cases_dist.head(8)
                    self._add_data_table(doc, "Cases % Distribution by Category", cases_dist_top, max_rows=8)
            
            # Lines Distribution Table  
            if 'lines_distribution' in category_analysis:
                lines_dist = category_analysis['lines_distribution']
                if not lines_dist.empty:
                    # Show top categories only
                    lines_dist_top = lines_dist.head(8)
                    self._add_data_table(doc, "Lines % Distribution by Category", lines_dist_top, max_rows=8)
            
            # Add key findings summary
            if 'slotting_insights' in category_analysis:
                insights = category_analysis['slotting_insights']
                if 'key_findings' in insights:
                    findings = insights['key_findings']
                    
                    findings_data = []
                    if findings.get('top_volume_category'):
                        findings_data.append(['Top Volume Category', findings['top_volume_category']])
                    if findings.get('top_velocity_category'):
                        findings_data.append(['Top Velocity Category', findings['top_velocity_category']])
                    if findings.get('critical_abc_fms_classes'):
                        critical_classes = ', '.join(findings['critical_abc_fms_classes'])
                        findings_data.append(['Critical ABC-FMS Classes', critical_classes])
                    
                    if findings_data:
                        findings_df = pd.DataFrame(findings_data, columns=['Finding', 'Value'])
                        self._add_data_table(doc, "Key Performance Findings", findings_df, max_rows=3)
            
            doc.add_page_break()
        
        # Manpower & FTE Analysis
        if 'manpower_analysis' in analysis_results:
            self._add_heading_with_style(doc, "Workforce Planning & FTE Analysis", level=1)
            
            manpower_analysis = analysis_results['manpower_analysis']
            
            # Generate AI insights for manpower analysis
            try:
                manpower_insights = self.llm_integration.generate_manpower_summary(
                    manpower_analysis
                )
                if manpower_insights and not manpower_insights.startswith("("):
                    self._add_ai_insight_section(doc, "Workforce Planning Insights", manpower_insights)
            except Exception as e:
                self.logger.error(f"Failed to generate manpower insights: {e}")
            
            # Add FTE requirements summary table
            summary_metrics = manpower_analysis.get('summary_metrics', {})
            if summary_metrics:
                fte_summary_data = []
                fte_summary_data.append(['Recommended Core FTE', f"{summary_metrics.get('recommended_core_fte', 0):.0f}"])
                fte_summary_data.append(['Peak FTE Requirement', f"{summary_metrics.get('peak_fte_requirement', 0):.0f}"])
                fte_summary_data.append(['Average FTE Requirement', f"{summary_metrics.get('average_fte_requirement', 0):.1f}"])
                fte_summary_data.append(['Flex Capacity Needed', f"{summary_metrics.get('flex_capacity_needed', 0):.0f}"])
                fte_summary_data.append(['Peak Days (%)', f"{summary_metrics.get('peak_days_percentage', 0):.1f}%"])
                
                fte_summary_df = pd.DataFrame(fte_summary_data, columns=['FTE Metric', 'Value'])
                self._add_data_table(doc, "FTE Requirements Summary", fte_summary_df, max_rows=5)
            
            # Add cost analysis table
            cost_analysis = manpower_analysis.get('cost_analysis', {})
            if cost_analysis:
                cost_data = []
                monthly_budget = cost_analysis.get('total_monthly_labor', 0)
                cost_per_case = cost_analysis.get('cost_per_case', 0)
                annual_budget = cost_analysis.get('annual_labor_budget', 0)
                
                cost_data.append(['Monthly Labor Budget', f"${monthly_budget:,.0f}"])
                cost_data.append(['Annual Labor Budget', f"${annual_budget:,.0f}"])
                cost_data.append(['Cost per Case', f"${cost_per_case:.2f}"])
                
                cost_df = pd.DataFrame(cost_data, columns=['Cost Metric', 'Value'])
                self._add_data_table(doc, "Labor Cost Analysis", cost_df, max_rows=3)
            
            # Add embedded charts if available
            charts_dir = self.charts_dir
            
            # FTE requirements timeline chart
            fte_timeline_chart = charts_dir / 'fte_requirements_timeline.png'
            if fte_timeline_chart.exists():
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(str(fte_timeline_chart), width=Inches(6.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    caption = doc.add_paragraph("Figure: Daily FTE Requirements Timeline with Peak Analysis")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                    
                except Exception as e:
                    self.logger.warning(f"Failed to embed FTE timeline chart: {e}")
                    placeholder = doc.add_paragraph("📊 Daily FTE Requirements Timeline Chart")
                    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    placeholder.style = 'Intense Quote'
            
            # Workforce planning dashboard chart
            workforce_dashboard_chart = charts_dir / 'workforce_planning_dashboard.png'
            if workforce_dashboard_chart.exists():
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(str(workforce_dashboard_chart), width=Inches(7.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    caption = doc.add_paragraph("Figure: Comprehensive Workforce Planning Dashboard")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                    
                except Exception as e:
                    self.logger.warning(f"Failed to embed workforce dashboard chart: {e}")
                    placeholder = doc.add_paragraph("📊 Workforce Planning Dashboard Chart")
                    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    placeholder.style = 'Intense Quote'
            
            # Add category labor distribution if available
            category_analysis = manpower_analysis.get('category_labor_analysis', {})
            category_breakdown = category_analysis.get('category_breakdown')
            if category_breakdown is not None and not category_breakdown.empty:
                category_labor_data = []
                for _, row in category_breakdown.iterrows():
                    category_labor_data.append([
                        row['Category'],
                        f"{row['FTE_Total_Required']:.1f}",
                        f"{row['FTE_Percentage']:.1f}%",
                        f"{row['Labor_Coefficient']:.2f}"
                    ])
                
                if category_labor_data:
                    category_df = pd.DataFrame(category_labor_data, 
                                             columns=['Category', 'FTE Required', 'FTE %', 'Complexity Factor'])
                    self._add_data_table(doc, "Category Labor Allocation", category_df, max_rows=5)
            
            # Add shift planning recommendations
            shift_planning = manpower_analysis.get('shift_planning', {})
            if shift_planning:
                core_staffing = shift_planning.get('core_staffing', {})
                peak_staffing = shift_planning.get('peak_staffing', {})
                
                if core_staffing and peak_staffing:
                    shift_data = []
                    shift_data.append(['Core Day Shift', f"{core_staffing.get('shift_1_day', 0):.0f} FTE"])
                    shift_data.append(['Core Evening Shift', f"{core_staffing.get('shift_2_evening', 0):.0f} FTE"])
                    shift_data.append(['Peak Day Shift', f"{peak_staffing.get('shift_1_day', 0):.0f} FTE"])
                    shift_data.append(['Peak Evening Shift', f"{peak_staffing.get('shift_2_evening', 0):.0f} FTE"])
                    
                    shift_df = pd.DataFrame(shift_data, columns=['Shift Configuration', 'Staffing Level'])
                    self._add_data_table(doc, "Recommended Shift Planning", shift_df, max_rows=4)
            
            doc.add_page_break()


def generate_word_report(analysis_results: Dict, 
                        output_dir: Path = REPORT_DIR,
                        filename: str = None) -> Path:
    """
    Convenience function to generate a Word report.
    
    Args:
        analysis_results: Dictionary containing analysis results
        output_dir: Output directory
        filename: Output filename
        
    Returns:
        Path to generated Word document
    """
    generator = WordReportGenerator(output_dir=output_dir)
    return generator.generate_word_report(analysis_results, filename)